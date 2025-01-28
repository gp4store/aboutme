from flask import Flask, render_template, request, send_file
from docx import Document
import os

app = Flask(__name__)

@app.route('/')
def home():
    return '''
        <h1>Generate and Download .docx File</h1>
        <form action="/generate-doc" method="post">
            <label for="content">Enter content for the Word document:</label><br>
            <textarea id="content" name="content" rows="10" cols="30"></textarea><br><br>
            <button type="submit">Generate and Download</button>
        </form>
    '''

@app.route('/generate-doc', methods=['POST'])
def generate_doc():
    # Get content from the form
    content = request.form.get('content', 'Default content for the Word document.')
    
    # Create a Word document
    doc = Document()
    doc.add_heading('Generated Document', level=1)
    doc.add_paragraph(content)
    
    # Save the document to a file
    filename = "generated_document.docx"
    filepath = os.path.join(os.getcwd(), filename)
    doc.save(filepath)
    
    # Serve the file for download
    return send_file(filepath, as_attachment=True, download_name=filename)

if __name__ == '__main__':
    app.run(debug=True)
